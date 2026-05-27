from __future__ import annotations

import io
import json
import logging
from datetime import datetime
from typing import List, Optional

from fastapi import (
    APIRouter,
    Depends,
    File,
    Form,
    HTTPException,
    Path as PathParam,
    Query,
    Request,
    UploadFile,
    status,
)
from fastapi.concurrency import run_in_threadpool
from fastapi.responses import Response, StreamingResponse
from sqlalchemy.exc import SQLAlchemyError
from sqlalchemy.orm import Session

from app.core.config import get_settings
from app.dependencies import get_current_user, get_db
from app.model.scan import Scan as ScanModel
from app.model.user import User
from app.schemas.scan import (
    DetectResponse,
    ScanAssets,
    ScanCreateResponse,
    ScanListItem,
    ScanOut,
)
from app.services import storage
from ml import detector as ml_detector, pipeline, scanner_enhance

router = APIRouter(prefix="/scan", tags=["scan"])
log = logging.getLogger(__name__)
_settings = get_settings()

_VALID_MODES = ("auto", "printed", "handwriting")
_VALID_ASSET_KINDS = ("raw", "enhanced", "pdf", "docx")
_VALID_ENGINES = ("pytesseract", "from_scratch")
_DOCX_MIME = (
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
)

def _validate_mode(mode: str) -> None:
    if mode not in _VALID_MODES:
        raise HTTPException(status_code=400, detail="invalid mode")

def _validate_enhance(mode: str) -> None:
    if mode not in scanner_enhance.MODES:
        raise HTTPException(
            status_code=400,
            detail=f"invalid enhance_mode; allowed: {','.join(scanner_enhance.MODES)}",
        )

def _validate_engine(engine: str) -> None:
    if engine not in _VALID_ENGINES:
        raise HTTPException(
            status_code=400,
            detail=f"invalid ocr_engine; allowed: {','.join(_VALID_ENGINES)}",
        )

def _parse_quad_override(raw: Optional[str]):
    if not raw:
        return None
    try:
        data = json.loads(raw)
    except Exception as e:
        raise HTTPException(status_code=400, detail="quad_override is not valid JSON") from e
    if not isinstance(data, list) or len(data) != 4:
        raise HTTPException(status_code=400, detail="quad_override must be 4 [x,y] pairs")
    try:
        import numpy as np

        arr = np.asarray(
            [[float(p[0]), float(p[1])] for p in data], dtype=np.float32
        )
    except Exception as e:
        raise HTTPException(status_code=400, detail="quad_override has bad shape") from e
    if arr.shape != (4, 2):
        raise HTTPException(status_code=400, detail="quad_override must be 4 [x,y] pairs")
    return arr

async def _read_capped(file: UploadFile, request: Request) -> bytes:
    cap = _settings.max_upload_bytes
    cl = request.headers.get("content-length")
    if cl and cl.isdigit() and int(cl) > cap:
        raise HTTPException(
            status_code=413,
            detail=f"upload too large (max {cap} bytes)",
        )
    chunks: list[bytes] = []
    total = 0
    while True:
        chunk = await file.read(1024 * 1024)
        if not chunk:
            break
        total += len(chunk)
        if total > cap:
            raise HTTPException(
                status_code=413,
                detail=f"upload too large (max {cap} bytes)",
            )
        chunks.append(chunk)
    return b"".join(chunks)

def _asset_url(scan_id: str, kind: str) -> str:
    return f"/scan/{scan_id}/asset/{kind}"

def _to_scan_out(scan: ScanModel) -> ScanOut:
    return ScanOut(
        id=scan.id,
        name=scan.name,
        created_at=scan.created_at,
        original_filename=scan.original_filename,
        bytes_size=scan.bytes_size,
        mode=scan.mode,
        enhance_mode=scan.enhance_mode,
        pipeline_path=scan.pipeline_path,
        ocr_engine=getattr(scan, "ocr_engine", None) or "pytesseract",
        language=scan.language,
        mean_conf=scan.mean_conf,
        document_detected=scan.document_detected,
        detection_score=scan.detection_score,
        handwriting_detected=scan.handwriting_detected,
        handwriting_confidence=scan.handwriting_confidence,
        confidence_warning=scan.confidence_warning,
        psm_used=scan.psm_used,
        text=scan.text,
        enhanced_mime=scan.enhanced_mime,
        assets=ScanAssets(
            raw=_asset_url(scan.id, "raw") if scan.raw_path else None,
            enhanced=_asset_url(scan.id, "enhanced") if scan.enhanced_path else None,
            pdf=_asset_url(scan.id, "pdf") if scan.pdf_path else None,
            docx=_asset_url(scan.id, "docx") if getattr(scan, "docx_path", None) else None,
        ),
    )

def _default_name(filename: Optional[str]) -> str:
    base = (filename or "").rsplit(".", 1)[0]
    if base.strip():
        return base.strip()[:120]
    now = datetime.utcnow()
    return f"Scan_{now:%Y%m%d_%H%M%S}"

def _get_owned_scan(scan_id: str, user: User, db: Session) -> ScanModel:
    scan = (
        db.query(ScanModel)
        .filter(ScanModel.id == scan_id, ScanModel.user_id == user.user_id)
        .first()
    )
    if scan is None:
        raise HTTPException(status_code=404, detail="scan not found")
    return scan

@router.post(
    "/extract",
    response_model=ScanCreateResponse,
    status_code=status.HTTP_201_CREATED,
)
async def extract(
    request: Request,
    file: UploadFile = File(...),
    lang: str = Form(default="eng"),
    spell_check: bool = Form(default=True),
    return_enhanced: bool = Form(default=True),
    mode: str = Form(default="auto"),
    enhance_mode: str = Form(default="color"),
    enhance_quality: int = Form(default=88),
    name: Optional[str] = Form(default=None),
    ocr_engine: str = Form(default="pytesseract"),
    quad_override: Optional[str] = Form(default=None),
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    _validate_mode(mode)
    _validate_enhance(enhance_mode)
    _validate_engine(ocr_engine)
    quad_arr = _parse_quad_override(quad_override)

    data = await _read_capped(file, request)
    if not data:
        raise HTTPException(status_code=400, detail="empty file")

    try:
        result = await run_in_threadpool(
            pipeline.run_from_bytes,
            data,
            lang,
            spell_check,
            False,
            mode,
            enhance_mode,
            enhance_quality,
            return_enhanced,
            False,
            ocr_engine,
            quad_arr,
        )
    except ValueError as e:
        raise HTTPException(status_code=400, detail=str(e))
    except HTTPException:
        raise
    except Exception:
        log.exception("OCR pipeline failed")
        raise HTTPException(status_code=500, detail="ocr failed")

    scan = ScanModel(
        user_id=current_user.user_id,
        name=name or _default_name(file.filename),
        original_filename=file.filename,
        bytes_size=len(data),
        mode=mode,
        enhance_mode=result.enhance_mode_used,
        pipeline_path=result.pipeline_path,
        ocr_engine=result.ocr_engine_used,
        mean_conf=result.mean_conf,
        document_detected=result.document_detected,
        detection_score=result.detection_score,
        handwriting_detected=result.handwriting_detected,
        handwriting_confidence=result.handwriting_confidence,
        confidence_warning=result.confidence_warning,
        psm_used=result.psm_used,
        language=result.language,
        text=result.text,
        enhanced_mime=result.enhanced_mime if result.enhanced_bytes else None,
    )
    db.add(scan)
    try:
        db.flush()
    except SQLAlchemyError:
        db.rollback()
        log.exception(
            "DB flush failed for new scan; numeric field types: "
            "mean_conf=%s detection_score=%s handwriting_confidence=%s psm_used=%s "
            "document_detected=%s handwriting_detected=%s",
            type(scan.mean_conf).__name__,
            type(scan.detection_score).__name__,
            type(scan.handwriting_confidence).__name__,
            type(scan.psm_used).__name__,
            type(scan.document_detected).__name__,
            type(scan.handwriting_detected).__name__,
        )
        raise HTTPException(status_code=500, detail="failed to persist scan")

    try:
        scan.raw_path = storage.save_bytes(
            current_user.user_id, scan.id, "raw", "image/jpeg", data
        )
        if result.enhanced_bytes:
            scan.enhanced_path = storage.save_bytes(
                current_user.user_id,
                scan.id,
                "enhanced",
                result.enhanced_mime,
                result.enhanced_bytes,
            )
    except Exception:
        log.exception("storage write failed for scan %s", scan.id)
        db.rollback()
        storage.delete_scan_dir(current_user.user_id, scan.id)
        raise HTTPException(status_code=500, detail="failed to persist scan assets")

    db.commit()
    db.refresh(scan)
    return _to_scan_out(scan)

@router.post("/preview")
async def preview(
    request: Request,
    file: UploadFile = File(...),
    enhance_mode: str = Form(default="color"),
    enhance_quality: int = Form(default=88),
    mode: str = Form(default="auto"),
    quad_override: Optional[str] = Form(default=None),
    current_user: User = Depends(get_current_user),
):
    _validate_mode(mode)
    _validate_enhance(enhance_mode)
    quad_arr = _parse_quad_override(quad_override)
    data = await _read_capped(file, request)
    if not data:
        raise HTTPException(status_code=400, detail="empty file")
    try:
        result = await run_in_threadpool(
            pipeline.run_from_bytes,
            data,
            "eng",
            False,
            False,
            mode,
            enhance_mode,
            enhance_quality,
            True,
            True,
            "pytesseract",
            quad_arr,
        )
    except ValueError as e:
        raise HTTPException(status_code=400, detail=str(e))
    except HTTPException:
        raise
    except Exception:
        log.exception("preview pipeline failed")
        raise HTTPException(status_code=500, detail="preview failed")

    if not result.enhanced_bytes:
        raise HTTPException(status_code=500, detail="failed to encode enhanced image")
    return Response(
        content=result.enhanced_bytes,
        media_type=result.enhanced_mime,
        headers={"X-Enhance-Mode": result.enhance_mode_used},
    )

@router.post("/detect", response_model=DetectResponse)
async def detect(
    request: Request,
    file: UploadFile = File(...),
    current_user: User = Depends(get_current_user),
):
    data = await _read_capped(file, request)
    if not data:
        raise HTTPException(status_code=400, detail="empty file")

    def _detect_sync(image_bytes: bytes) -> tuple[bool, float, Optional[list[list[float]]], int, int]:
        img = pipeline._decode(image_bytes)
        det = ml_detector.detect_document(img)
        h, w = img.shape[:2]
        quad: Optional[list[list[float]]] = None
        if det.document_detected and det.quad is not None:
            quad = [[float(x), float(y)] for (x, y) in det.quad.reshape(4, 2)]
        return det.document_detected, float(det.score), quad, int(w), int(h)

    try:
        document_detected, score, quad, w, h = await run_in_threadpool(_detect_sync, data)
    except ValueError as e:
        raise HTTPException(status_code=400, detail=str(e))
    except Exception:
        log.exception("detect failed")
        raise HTTPException(status_code=500, detail="detect failed")

    return DetectResponse(
        document_detected=document_detected,
        score=score,
        quad=quad,
        image_width=w,
        image_height=h,
    )

@router.post("/{scan_id}/reprocess", response_model=ScanOut)
async def reprocess(
    scan_id: str = PathParam(...),
    mode: str = Form(...),
    enhance_mode: Optional[str] = Form(default=None),
    enhance_quality: int = Form(default=88),
    lang: str = Form(default="eng"),
    spell_check: bool = Form(default=True),
    ocr_engine: str = Form(default="pytesseract"),
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    _validate_mode(mode)
    _validate_engine(ocr_engine)
    target_enhance = enhance_mode or "color"
    _validate_enhance(target_enhance)

    scan = _get_owned_scan(scan_id, current_user, db)
    if not scan.raw_path:
        raise HTTPException(status_code=400, detail="scan has no raw asset")

    try:
        raw = storage.read_bytes(scan.raw_path)
    except FileNotFoundError:
        raise HTTPException(status_code=404, detail="raw asset missing on disk")
    except PermissionError:
        raise HTTPException(status_code=403, detail="forbidden")

    try:
        result = await run_in_threadpool(
            pipeline.run_from_bytes,
            raw,
            lang,
            spell_check,
            False,
            mode,
            target_enhance,
            enhance_quality,
            True,
            False,
            ocr_engine,
            None,
        )
    except ValueError as e:
        raise HTTPException(status_code=400, detail=str(e))
    except HTTPException:
        raise
    except Exception:
        log.exception("reprocess pipeline failed for scan %s", scan_id)
        raise HTTPException(status_code=500, detail="reprocess failed")

    if result.enhanced_bytes:
        try:
            scan.enhanced_path = storage.save_bytes(
                current_user.user_id,
                scan.id,
                "enhanced",
                result.enhanced_mime,
                result.enhanced_bytes,
            )
            scan.enhanced_mime = result.enhanced_mime
        except Exception:
            log.exception("storage write failed during reprocess for scan %s", scan_id)
            db.rollback()
            raise HTTPException(
                status_code=500, detail="failed to persist enhanced asset"
            )

    scan.mode = mode
    scan.enhance_mode = result.enhance_mode_used
    scan.pipeline_path = result.pipeline_path
    scan.ocr_engine = result.ocr_engine_used
    scan.mean_conf = result.mean_conf
    scan.handwriting_detected = result.handwriting_detected
    scan.handwriting_confidence = result.handwriting_confidence
    scan.confidence_warning = result.confidence_warning
    scan.psm_used = result.psm_used
    scan.language = result.language
    scan.text = result.text

    db.commit()
    db.refresh(scan)
    return _to_scan_out(scan)

@router.post("/{scan_id}/pdf", response_model=ScanOut)
async def attach_pdf_to_scan(
    request: Request,
    scan_id: str = PathParam(...),
    enhance_mode: str = Form(default="color"),
    enhance_quality: int = Form(default=88),
    mode: str = Form(default="auto"),
    searchable: bool = Form(default=True),
    lang: str = Form(default="eng"),
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    _validate_mode(mode)
    _validate_enhance(enhance_mode)
    scan = _get_owned_scan(scan_id, current_user, db)
    if not scan.raw_path:
        raise HTTPException(status_code=400, detail="scan has no raw asset")

    try:
        raw = storage.read_bytes(scan.raw_path)
    except Exception:
        log.exception("failed to read raw asset for scan %s", scan_id)
        raise HTTPException(status_code=500, detail="failed to read scan asset")

    try:
        pdf_bytes = await run_in_threadpool(
            _build_pdf,
            [raw],
            mode,
            enhance_mode,
            enhance_quality,
            searchable,
            lang,
        )
    except HTTPException:
        raise
    except Exception:
        log.exception("PDF generation failed for scan %s", scan_id)
        raise HTTPException(status_code=500, detail="pdf failed")

    scan.pdf_path = storage.save_bytes(
        current_user.user_id, scan.id, "pdf", "application/pdf", pdf_bytes
    )
    db.commit()
    db.refresh(scan)
    return _to_scan_out(scan)


@router.post("/{scan_id}/docx", response_model=ScanOut)
async def attach_docx_to_scan(
    scan_id: str = PathParam(...),
    include_image: bool = Form(default=True),
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    scan = _get_owned_scan(scan_id, current_user, db)
    if not scan.text and not scan.enhanced_path:
        raise HTTPException(
            status_code=400, detail="scan has no text or enhanced asset to export"
        )

    image_bytes: Optional[bytes] = None
    if include_image and scan.enhanced_path:
        try:
            image_bytes = storage.read_bytes(scan.enhanced_path)
        except Exception:
            log.warning("could not read enhanced asset for docx; embedding skipped")
            image_bytes = None

    try:
        docx_bytes = await run_in_threadpool(
            _build_docx,
            scan.name or "Scan",
            scan.text or "",
            image_bytes,
        )
    except HTTPException:
        raise
    except Exception:
        log.exception("DOCX generation failed for scan %s", scan_id)
        raise HTTPException(status_code=500, detail="docx failed")

    scan.docx_path = storage.save_bytes(
        current_user.user_id, scan.id, "docx", _DOCX_MIME, docx_bytes
    )
    db.commit()
    db.refresh(scan)
    return _to_scan_out(scan)

@router.post("/pdf")
async def standalone_pdf(
    request: Request,
    files: List[UploadFile] = File(...),
    enhance_mode: str = Form(default="color"),
    enhance_quality: int = Form(default=88),
    mode: str = Form(default="auto"),
    searchable: bool = Form(default=True),
    lang: str = Form(default="eng"),
    current_user: User = Depends(get_current_user),
):
    _validate_mode(mode)
    _validate_enhance(enhance_mode)
    if not files:
        raise HTTPException(status_code=400, detail="no files")
    if len(files) > _settings.max_pdf_pages:
        raise HTTPException(
            status_code=400,
            detail=f"too many pages (max {_settings.max_pdf_pages})",
        )

    pages: list[bytes] = []
    for f in files:
        data = await _read_capped(f, request)
        if data:
            pages.append(data)
    if not pages:
        raise HTTPException(status_code=400, detail="no usable pages")

    try:
        pdf_bytes = await run_in_threadpool(
            _build_pdf,
            pages,
            mode,
            enhance_mode,
            enhance_quality,
            searchable,
            lang,
        )
    except HTTPException:
        raise
    except Exception:
        log.exception("PDF generation failed")
        raise HTTPException(status_code=500, detail="pdf failed")
    return Response(content=pdf_bytes, media_type="application/pdf")

def _build_pdf(
    page_bytes_list: list[bytes],
    mode: str,
    enhance_mode: str,
    enhance_quality: int,
    searchable: bool,
    lang: str,
) -> bytes:
    try:
        from PIL import Image
    except Exception as e:
        raise HTTPException(status_code=500, detail="Pillow not installed") from e

    page_results = []
    enhanced_pages: list[bytes] = []
    for raw in page_bytes_list:
        result = pipeline.run_from_bytes(
            raw,
            lang,
            True,
            False,
            mode,
            enhance_mode,
            enhance_quality,
            True,
            not searchable,
        )
        if not result.enhanced_bytes:
            continue
        page_results.append(result)
        enhanced_pages.append(result.enhanced_bytes)
    if not enhanced_pages:
        raise HTTPException(status_code=400, detail="no usable pages")

    use_searchable = searchable and all(
        p.pipeline_path == "printed" and p.mean_conf >= 60 and p.enhance_mode_used != "bw"
        for p in page_results
    )

    if use_searchable:
        try:
            import pytesseract

            chunks: list[bytes] = []
            for raw in enhanced_pages:
                im = Image.open(io.BytesIO(raw)).convert("RGB")
                chunk = pytesseract.image_to_pdf_or_hocr(
                    im, lang=lang, extension="pdf", config="--oem 3 --psm 6"
                )
                chunks.append(chunk if isinstance(chunk, (bytes, bytearray)) else bytes(chunk, "latin-1"))
            if len(chunks) == 1:
                return chunks[0]
            try:
                from pypdf import PdfReader, PdfWriter

                writer = PdfWriter()
                for c in chunks:
                    reader = PdfReader(io.BytesIO(c))
                    for page in reader.pages:
                        writer.add_page(page)
                out = io.BytesIO()
                writer.write(out)
                return out.getvalue()
            except Exception:
                log.warning("pypdf merge failed; falling back to image-only PDF")
        except Exception:
            log.exception("searchable PDF generation failed; falling back")

    images = [Image.open(io.BytesIO(raw)).convert("RGB") for raw in enhanced_pages]
    out = io.BytesIO()
    if len(images) == 1:
        images[0].save(out, format="PDF")
    else:
        images[0].save(out, format="PDF", save_all=True, append_images=images[1:])
    return out.getvalue()


def _build_docx(title: str, text: str, image_bytes: Optional[bytes]) -> bytes:
    """Render a minimal .docx with optional embedded enhanced image + text."""
    try:
        from docx import Document
        from docx.shared import Inches, Pt
    except Exception as e:
        raise HTTPException(status_code=500, detail="python-docx not installed") from e

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    h = doc.add_heading(title or "Scan", level=1)
    for run in h.runs:
        run.font.size = Pt(18)

    if image_bytes:
        try:
            doc.add_picture(io.BytesIO(image_bytes), width=Inches(6.0))
        except Exception:
            log.warning("docx: failed to embed image; continuing with text only")

    if text.strip():
        for paragraph in text.split("\n\n"):
            doc.add_paragraph(paragraph)
    else:
        doc.add_paragraph("(no extracted text)")

    out = io.BytesIO()
    doc.save(out)
    return out.getvalue()

@router.get("/history", response_model=list[ScanListItem])
def history(
    limit: int = Query(default=50, ge=1, le=200),
    offset: int = Query(default=0, ge=0),
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    rows = (
        db.query(ScanModel)
        .filter(ScanModel.user_id == current_user.user_id)
        .order_by(ScanModel.created_at.desc())
        .offset(offset)
        .limit(limit)
        .all()
    )
    return [
        ScanListItem(
            id=r.id,
            name=r.name,
            created_at=r.created_at,
            bytes_size=r.bytes_size,
            pipeline_path=r.pipeline_path,
            enhance_mode=r.enhance_mode,
            ocr_engine=getattr(r, "ocr_engine", None) or "pytesseract",
            handwriting_detected=r.handwriting_detected,
            mean_conf=r.mean_conf,
            has_pdf=bool(r.pdf_path),
            has_enhanced=bool(r.enhanced_path),
            has_docx=bool(getattr(r, "docx_path", None)),
        )
        for r in rows
    ]

@router.get("/{scan_id}", response_model=ScanOut)
def get_scan(
    scan_id: str,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    scan = _get_owned_scan(scan_id, current_user, db)
    return _to_scan_out(scan)

@router.delete("/{scan_id}", status_code=status.HTTP_204_NO_CONTENT)
def delete_scan(
    scan_id: str,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    scan = _get_owned_scan(scan_id, current_user, db)
    db.delete(scan)
    db.commit()
    storage.delete_scan_dir(current_user.user_id, scan_id)
    return Response(status_code=status.HTTP_204_NO_CONTENT)

@router.get("/{scan_id}/asset/{kind}")
def get_asset(
    scan_id: str,
    kind: str,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    if kind not in _VALID_ASSET_KINDS:
        raise HTTPException(status_code=400, detail="invalid asset kind")
    scan = _get_owned_scan(scan_id, current_user, db)

    if kind == "raw":
        path = scan.raw_path
        media_type = "image/jpeg"
    elif kind == "enhanced":
        path = scan.enhanced_path
        media_type = scan.enhanced_mime or "image/jpeg"
    elif kind == "docx":
        path = getattr(scan, "docx_path", None)
        media_type = _DOCX_MIME
    else:
        path = scan.pdf_path
        media_type = "application/pdf"

    if not path:
        raise HTTPException(status_code=404, detail=f"asset {kind} not available")
    try:
        data = storage.read_bytes(path)
    except FileNotFoundError:
        raise HTTPException(status_code=404, detail="asset missing on disk")
    except PermissionError:
        log.error("path traversal blocked for scan %s kind %s", scan_id, kind)
        raise HTTPException(status_code=403, detail="forbidden")
    return StreamingResponse(io.BytesIO(data), media_type=media_type)
